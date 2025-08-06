import os
import logging
from io import BytesIO
import hashlib
from typing import List, Dict, Any

import google.generativeai as genai
# from pinecone.data.index import Index
from pinecone import Pinecone
from pinecone import PineconeException
from pypdf import PdfReader
from sentence_transformers import CrossEncoder
from typing import Optional
from docx import Document  
from typing import TYPE_CHECKING
if TYPE_CHECKING:
    from pinecone import Index

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
cross_encoder_model: Optional[CrossEncoder] = None

def detect_file_type(content: bytes) -> str:
    """
    Detect file type based on content headers.
    """
    if content.startswith(b'%PDF'):
        return 'pdf'
    elif content.startswith(b'PK') and b'word/' in content[:1000]:
        return 'docx'
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content."""
    try:
        reader = PdfReader(BytesIO(content))
        if len(reader.pages) == 0:
            raise ValueError("PDF contains no pages")
        
        full_text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                full_text += page_text
                logger.debug(f"Extracted text from PDF page {page_num + 1}")
            except Exception as e:
                logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {e}")
                continue
        
        return full_text
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content."""
    try:
        doc = Document(BytesIO(content))
        full_text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += cell.text + " "
                full_text += "\n"
        
        logger.info(f"Extracted text from DOCX: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        return full_text
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")

def extract_text_from_document(content: bytes, file_type: str = None) -> str:
    """
    Extract text from document based on file type.
    """
    if not content:
        raise ValueError("Document content cannot be empty")
    
    # Auto-detect file type if not provided
    if not file_type:
        file_type = detect_file_type(content)
    
    logger.info(f"Processing {file_type.upper()} document")
    
    if file_type == 'pdf':
        text = extract_text_from_pdf(content)
    elif file_type == 'docx':
        text = extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    if not text.strip():
        raise ValueError(f"No text could be extracted from the {file_type.upper()} document")
    
    return text

def smart_chunk_text(text: str, max_chunk_size: int = 1500, overlap: int = 200) -> List[str]:
    """
    Intelligently chunk text by sentences to avoid breaking mid-sentence.
    """
    if not text or not text.strip():
        return []
    
    # Split by sentences (basic approach)
    sentences = text.replace('!', '.').replace('?', '.').split('.')
    sentences = [s.strip() + '.' for s in sentences if s.strip()]
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would exceed the limit
        if len(current_chunk) + len(sentence) > max_chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                # Start new chunk with overlap from previous chunk
                words = current_chunk.split()
                overlap_text = ' '.join(words[-overlap//10:]) if len(words) > overlap//10 else ""
                current_chunk = overlap_text + " " + sentence if overlap_text else sentence
            else:
                # Single sentence is too long, force split
                chunks.append(sentence[:max_chunk_size])
                current_chunk = sentence[max_chunk_size:]
        else:
            current_chunk += " " + sentence if current_chunk else sentence
    
    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return [chunk for chunk in chunks if len(chunk.strip()) > 50]  # Filter tiny chunks

def process_document(index: Any, document_content: bytes, namespace: str, file_type: str = None):
    """
    Processes a document (PDF or DOCX) by extracting text, chunking, embedding, and upserting into Pinecone.
    """
    # Input validation
    if not index:
        raise ValueError("Index cannot be None")
    if not document_content:
        raise ValueError("Document content cannot be empty")
    if not namespace or not namespace.strip():
        raise ValueError("Namespace cannot be empty")
    
    try:
        # 1. Extract Text from Document
        logger.info(f"Starting document processing for namespace: {namespace}")
        
        full_text = extract_text_from_document(document_content, file_type)
        
        # Clean the text
        full_text = full_text.replace('\n', ' ').replace('\r', ' ')
        full_text = ' '.join(full_text.split())  # Remove extra whitespace
        
        logger.info(f"Extracted {len(full_text)} characters from document")
        
        # 2. Chunk the Text
        text_chunks = smart_chunk_text(full_text, max_chunk_size=1500)
        
        if not text_chunks:
            raise ValueError("Failed to create any text chunks from the document")
        
        logger.info(f"Created {len(text_chunks)} text chunks")
        
        # 3. Embed Chunks
        embedding_model = "models/text-embedding-004"
        
        try:
            embedding_result = genai.embed_content(
                model=embedding_model,
                content=text_chunks,
                task_type="RETRIEVAL_DOCUMENT",
                title="Document Chunks"
            )
            
            if not embedding_result or 'embedding' not in embedding_result:
                raise ValueError("Failed to generate embeddings")
                
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise ValueError(f"Failed to generate embeddings: {e}")
        
        # 4. Prepare vectors for upsert
        vectors_to_upsert = []
        for i, (chunk, embedding) in enumerate(zip(text_chunks, embedding_result['embedding'])):
            # Create unique ID using content hash + index
            chunk_hash = hashlib.md5(chunk.encode()).hexdigest()[:8]
            vector_id = f"{namespace}{chunk_hash}{i}"
            
            vectors_to_upsert.append({
                "id": vector_id,
                "values": embedding,
                "metadata": {
                    "text": chunk,
                    "chunk_index": i,
                    "namespace": namespace,
                    "char_count": len(chunk),
                    "file_type": file_type or "unknown"
                }
            })
        
        # 5. Upsert to Pinecone in batches
        batch_size = 50
        successful_upserts = 0
        
        for i in range(0, len(vectors_to_upsert), batch_size):
            batch = vectors_to_upsert[i:i + batch_size]
            try:
                index.upsert(vectors=batch, namespace=namespace)
                successful_upserts += len(batch)
                logger.info(f"Upserted batch {i//batch_size + 1}, vectors: {len(batch)}")
            except PineconeException as e:
                logger.error(f"Failed to upsert batch {i//batch_size + 1}: {e}")
                raise
        
        logger.info(f"Successfully processed and upserted {successful_upserts} chunks into namespace '{namespace}'")
        
    except ValueError:
        # Re-raise validation errors as-is
        raise
    except PineconeException as e:
        logger.error(f"Pinecone error during document processing: {e}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error during document processing: {e}")
        raise ValueError(f"Document processing failed: {e}")

def namespace_exists(index: Any, namespace: str) -> bool:
    """
    Checks if a namespace exists in the Pinecone index by querying its stats.
    """
    if not index or not namespace:
        raise ValueError("Index and namespace cannot be None or empty")
    
    try:
        stats = index.describe_index_stats()
        exists = namespace in stats.get('namespaces', {})
        logger.info(f"Namespace '{namespace}' exists: {exists}")
        return exists
    except PineconeException as e:
        logger.error(f"Pinecone API error checking namespace: {e}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error checking namespace: {e}")
        raise

def generate_verified_answer(context: str, question: str) -> str:
    """Generates a high-fidelity, concise answer using Gemini 2.5 Flash and an enhanced, minimalistic prompt."""
    logger.info("Generating final verified answer with Gemini 2.5 Flash...")

    # The enhanced prompt for a concise, minimalistic answer
    prompt = f"""You are a highly precise and minimalistic AI assistant. Your sole purpose is to answer a user's query based **EXCLUSIVELY** on the provided text context.

Follow these strict output rules:
- Provide a concise answer in 1-2 lines.
- Provide the most relevant sentences from the context that support the answer, also in 1-2 lines.
- Do not add any extra commentary, headers, or information. The output should be only the answer and the context.

**Instructions:**
1.  **Identify:** Scrutinize the provided context and find the exact sentences and data points relevant to the user's query.
2.  **Synthesize:** Condense the identified information into a direct, 1-2 line answer.
3.  **Verify:** Ensure every fact in your answer is present in the context and that the answer is completely devoid of unasked-for information.

---
*Provided Context:*
{context}
---
*User Query:*
{question}
---

*Final Output:*
"""
    
    try:
        # Use the Gemini 2.5 Flash model as requested
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        response = model.generate_content(
            prompt,
            generation_config={
                'temperature': 0.0 # Set to 0.0 for maximum factuality
            }
        )
        return response.text.strip()
    except Exception as e:
        logger.error(f"Gemini generation failed: {e}")
        return "There was an error generating the final answer."

def generate_hypothetical_document(question: str) -> str:
    """Uses a fast LLM to generate a hypothetical answer to a question (HyDE)."""
    logger.info(f"Generating HyDE document for question: '{question}'")
    try:
        model = genai.GenerativeModel('gemini-2.5-flash')
        prompt = f"""Write a short, ideal answer paragraph for the following user question. This paragraph should be written in the formal style of a policy document and will be used to find similar-sounding passages in a vector search.

        Question: "{question}"
        
        Ideal Answer Paragraph:
        """
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        logger.error(f"HyDE generation failed: {e}")
        return question # Fallback to using the original question

def get_cross_encoder_model() -> CrossEncoder:
    """
    Lazily loads the CrossEncoder model on the first call and caches it.
    This is the singleton pattern.
    """
    global cross_encoder_model
    
    if cross_encoder_model is None:
        # This block only runs ONCE, during the first API request
        logger.info("Lazy loading CrossEncoder model for the first time...")
        
        model_path = 'cross-encoder/ms-marco-TinyBERT-L-2-v2'
        cross_encoder_model = CrossEncoder(model_path)
        logger.info("CrossEncoder model loaded successfully.")
        
    return cross_encoder_model

def rerank_with_cross_encoder(question: str, chunks: list) -> list:
    """Reranks retrieved chunks using the lazily-loaded Cross-Encoder model."""
    logger.info(f"Re-ranking {len(chunks)} chunks...")
    if not chunks:
        return []
    
    # 2. Call the getter function to get the model instance
    model = get_cross_encoder_model()
    
    # The rest of the function remains the same
    pairs = [[question, chunk['metadata']['text']] for chunk in chunks]
    scores = model.predict(pairs)
    
    for i in range(len(chunks)):
        chunks[i]['rerank_score'] = scores[i]
            
    return sorted(chunks, key=lambda x: x['rerank_score'], reverse=True)

def get_answer_for_question(index: 'Index', question: str, namespace: str) -> str:
    logger.info(f"Starting advanced RAG for question: '{question}'")

    # === PHASE 2: HYDE RETRIEVAL ===
    hypothetical_doc = generate_hypothetical_document(question)
    
    # Embed the HYPOTHETICAL document for the search
    search_embedding = genai.embed_content(
        model="models/text-embedding-004", 
        content=hypothetical_doc, 
        task_type="RETRIEVAL_DOCUMENT" # Use this type as the hypothetical doc is like a document
    )['embedding']
    
    # Retrieve a broad set of 20 candidates for re-ranking
    retrieved_chunks = index.query(
        namespace=namespace, 
        vector=search_embedding, 
        top_k=20, 
        include_metadata=True
    ).get('matches', [])

    if not retrieved_chunks:
        return "Could not find any relevant information in the document to answer the question."

    # === PHASE 3: CROSS-ENCODER RE-RANKING ===
    # Re-rank the 20 chunks against the ORIGINAL question for precision
    reranked_chunks = rerank_with_cross_encoder(question, retrieved_chunks)
    
    # Select the top 5 most relevant chunks for the final context
    top_k_reranked = reranked_chunks[:5]
    
    # Assemble the final context from the best chunks
    final_context = "\n\n".join([chunk['metadata']['text'] for chunk in top_k_reranked])

    # === PHASE 1: ADVANCED GENERATION ===
    # Generate the final, verified answer using the best context and GPT-4
    final_answer = generate_verified_answer(final_context, question)
    
    return final_answer